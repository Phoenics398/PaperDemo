from pdb import set_trace as sc

import os
import docx
import argparse

from docx import RT
from docx import table
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.shared import Inches, Cm, Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from sympy import simplify, symbols
# from latex2word import LatexToWordElement


__all__ = [
    'PaperDocx', 
]


class PaperDocx(object):
    """"""

    def __init__(self, path=''):
        """"""
        if path == '':
            self.doc = docx.Document()
        else:
            self.doc = docx.Document(path)

    def setProperties(self, obj, prop=None):
        """
        页面、段落、文字、页眉、页脚属性设置
        """
        tmp = vars(prop)
        for kx, vx in tmp.items():
            if vx != None and kx == 'settings_odd_and_even_pages_header_footer': 
                self.doc.settings.odd_and_even_pages_header_footer = vx

            elif vx != None and kx == 'page_width': 
                obj.page_width = vx
            elif vx != None and kx == 'page_height': 
                obj.page_height = vx
            elif vx != None and kx == 'left_margin': 
                obj.left_margin = vx
            elif vx != None and kx == 'right_margin': 
                obj.right_margin = vx
            elif vx != None and kx == 'top_margin': 
                obj.top_margin = vx
            elif vx != None and kx == 'bottom_margin': 
                obj.bottom_margin = vx
            elif vx != None and kx == 'gutter': 
                obj.gutter = vx
            elif vx != None and kx == 'header_distance': 
                obj.header_distance = vx
            elif vx != None and kx == 'footer_distance': 
                obj.footer_distance = vx
            elif vx != None and kx == 'orientation': 
                obj.orientation = vx
            elif vx != None and kx == 'start_type': 
                obj.start_type = vx
            elif vx != None and kx == 'different_first_page_header_footer': 
                obj.different_first_page_header_footer = vx
            elif vx != None and kx == 'header_is_linked_to_previous': 
                obj.header.is_linked_to_previous = vx
            elif vx != None and kx == 'footer_is_linked_to_previous': 
                obj.footer.is_linked_to_previous = vx
            elif vx != None and kx == 'alignment': 
                obj.alignment = vx
            elif vx != None and kx == 'paragraph_format_first_line_indent': 
                obj.paragraph_format.first_line_indent = vx
            elif vx != None and kx == 'paragraph_format_left_indent': 
                obj.paragraph_format.left_indent = vx
            elif vx != None and kx == 'paragraph_format_right_indent': 
                obj.paragraph_format.right_indent = vx
            elif vx != None and kx == 'paragraph_format_space_before': 
                obj.paragraph_format.space_before = vx
            elif vx != None and kx == 'paragraph_format_space_after': 
                obj.paragraph_format.space_after = vx
            elif vx != None and kx == 'paragraph_format_line_spacing_rule': 
                obj.paragraph_format.line_spacing_rule = vx
            elif vx != None and kx == 'paragraph_format_line_spacing': 
                obj.paragraph_format.line_spacing = vx
            elif vx != None and kx == 'paragraph_format_widow_control': 
                obj.paragraph_format.widow_control = vx
            elif vx != None and kx == 'paragraph_format_keep_together': 
                obj.paragraph_format.keep_together = vx
            elif vx != None and kx == 'paragraph_format_keep_with_next': 
                obj.paragraph_format.keep_with_next = vx
            elif vx != None and kx == 'paragraph_format_page_break_before': 
                obj.paragraph_format.page_break_before = vx

        if 'docx.text.paragraph.Paragraph' in str(type(obj)):
            for rx in obj.runs:
                for kx, vx in tmp.items():
                    if vx != None and kx == 'font_name':
                        rx.font.name = vx 
                    elif vx != None and kx == 'font_name_cn': 
                        rx.element.rPr.rFonts.set(qn('w:eastAsia'), vx)
                    elif vx != None and kx == 'font_size': 
                        rx.font.size = vx 
                    elif vx != None and kx == 'font_bold': 
                        rx.font.bold = vx 
                    elif vx != None and kx == 'font_italic': 
                        rx.font.italic = vx 
                    elif vx != None and kx == 'font_shadow': 
                        rx.font.shadow = vx 
                    elif vx != None and kx == 'font_outline': 
                        rx.font.outline = vx 
                    elif vx != None and kx == 'font_emboss': 
                        rx.font.emboss = vx 
                    elif vx != None and kx == 'font_rtl': 
                        rx.font.rtl = vx 
                    elif vx != None and kx == 'font_underline': 
                        rx.font.underline = vx 
                    elif vx != None and kx == 'font_math': 
                        rx.font.math = vx 
                    elif vx != None and kx == 'font_strike': 
                        rx.font.strike = vx 
                    elif vx != None and kx == 'font_double_strike': 
                        rx.font.double_strike = vx 
                    elif vx != None and kx == 'font_highlight_color': 
                        rx.font.highlight_color = vx 
                    elif vx != None and kx == 'font_superscript': 
                        rx.font.superscript = vx 
                    elif vx != None and kx == 'font_subscript': 
                        rx.font.subscript = vx 
                    elif vx != None and kx == 'font_imprint': 
                        rx.font.imprint = vx 
                    elif vx != None and kx == 'font_hidden': 
                        rx.font.hidden = vx 
                    elif vx != None and kx == 'font_no_proof': 
                        rx.font.no_proof = vx 
                    elif vx != None and kx == 'font_all_caps': 
                        rx.font.all_caps = vx 
                    elif vx != None and kx == 'font_small_caps': 
                        rx.font.small_caps = vx 
                    elif vx != None and kx == 'font_snap_to_grid': 
                        rx.font.snap_to_grid = vx 
                    elif vx != None and kx == 'font_spec_vanish': 
                        rx.font.spec_vanish = vx 
                    elif vx != None and kx == 'font_web_hidden': 
                        rx.font.web_hidden = vx 

        return self

    def setSection(self, prop=None):
        """
        章节设置
        """
        section = self.doc.sections[0]
        self.setProperties(section, prop)
        return self

    def addStyle(self, prop=None):
        """
        添加式样
        """
        style = self.doc.styles.add_style(name='phx', style_type=1, builtin=False)
        self.setProperties(style, prop)
        return style
    
    def addPage(self, section=None, prop=None, page_before='', page_after=''):
        """
        添加页码
        """
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        if section == None:
            section = self.doc.sections[0]
        tmp = section.footer.add_paragraph()
        tmp.add_run(page_before)
        tmp.add_run()._r.append(fldChar1)
        tmp.add_run()._r.append(instrText)
        tmp.add_run()._r.append(fldChar2)
        tmp.add_run(page_after)
        if prop != None:
            self.setProperties(tmp, prop)
        return self

    def addTitle(self, title, prop=None):
        """
        添加标题
        """
        tmp = self.doc.add_paragraph(title.title)
        if title.project_tag != '':
            tag = tmp.add_run(title.project_tag)
        if prop != None:
            self.setProperties(tmp, prop)
        tag.font.superscript = True
        return tmp
    
    def addTitleFooter(self, footer, section=None, prop=None):
        """
        添加标题注释
        """
        if section == None:
            section = self.doc.sections[0]
        tmp = section.footer.add_paragraph(footer)
        if prop != None:
            self.setProperties(tmp, prop)

    def exportPaper(self, path):
        """"""
        self.doc.save(path)

